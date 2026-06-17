# app/parsers/docx_parser.py
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.table import Table

from app.parsers.base_parser import BaseDocumentParser


def _clean_docx_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"([\u4e00-\u9fff])([A-Za-z0-9])", r"\1 \2", text)
    text = re.sub(r"([A-Za-z0-9])([\u4e00-\u9fff])", r"\1 \2", text)
    return text.strip()


def _table_to_text(table: Table) -> str:
    rows: list[str] = []

    for row in table.rows:
        cells = [_clean_docx_text(cell.text) for cell in row.cells]
        cells = [cell for cell in cells if cell]
        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows).strip()


class DocxDocumentParser(BaseDocumentParser):
    def parse(self, path: Path) -> str:
        doc = Document(str(path))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = _clean_docx_text(paragraph.text)
            if text:
                parts.append(text)

        for table in doc.tables:
            text = _table_to_text(table)
            if text:
                parts.append(text)

        return "\n\n".join(parts).strip()